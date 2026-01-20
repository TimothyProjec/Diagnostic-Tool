"""
Diagnosis generation using GPT-4o-mini with custom medical report format
"""

from openai import OpenAI
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import io

@st.cache_resource
def get_openai_client():
    """Initialize OpenAI client"""
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def generate_diagnosis_from_transcript(transcript):
    """
    Generate medical diagnosis from consultation transcript using custom format
    
    Args:
        transcript (str): Transcribed consultation text
        
    Returns:
        str: Formatted medical report text
    """
    try:
        client = get_openai_client()
        
        system_prompt = """You are an expert medical scribe.
Generate formal medical reports in the exact format provided by the user.
Follow the template precisely and maintain professional medical documentation standards."""
        
        user_prompt = f"""Based on this transcription, create a formal medical report in the following format. 

If you did not get the information from the transcription, leave it blank. 

For each diagnosis give a reason in bracket and below diagnosis section give provisional diagnosis in the same manner with reasoning.

NO need of ** for headings.

For tables, make proper tables with lines and columns.

TRANSCRIPTION:
{transcript}

---

FORMAT TO FOLLOW:

Name:
Age :      DOB:   /   /
Language:          Ethnic:           Address:      Mobile No:
Patient UHID:   WHBG
Date of Initial Examination:
Referred by:                  Primary Physician:   Dr.
DOA                  DOD      Drs.   
DOA                   DOD      Dr.


ADR & ALLERGIES:



DIAGNOSIS: (give appropriate number of diagnoses with reasoning minimum 15 and maximum 15 as well also give provisional diagnosis below diagnosis section with reason after going through all the data sources)

1.




SPECIAL RISKS :
1.


HISTORY:




Past & Rx History:




CURRENT MEDICATIONS:

MEDICATION  DOSE  TIME
  AM  Noon  PM  HS







PHYSICAL EXAMINATION:

General Status:                                   Handedness:
Weight:      Kg.   Height:      cm.     Temp:      deg F      BMI:
BP:            sitting;          lying;           standing   mmHg
Pulse:      /min         Peripheral pulses:
Skin:
CVS:                  RS:                ABD:             Genitalia:

Head Circumference:       cm.

Cognitive Functions:     Educational level:

Language functions:
Memory recall items:        /5
MMSE:     / 30.      Addenbrooke's:
Hamilton's Depression Scale:

Eyes:
Visual Acuity:                          Fields:
EOM:                                       Pupils:
Optic Fundi:
ENT:             Teeth:            Swallowing:       Cough:          Speech:

Facial:
Hearing:            Other  Cr. Nerves:

Urinary Bladder:                      Nocturia:        / night
Gait:
Spine:                  Hips:                   Knees:              Ankles:

SLRT:   R:         L:           degrees elevation

MUSCLE POWER ASSESSMENT
Date  RUL RLL LUL LLL General
  Prox  Dist  Prox  Dist  Prox  Dist  Prox  Dist

ACTIVITIES OF DAILY LIVING (ADL) / Normal = 5, Limp Walking = 3, Bedridden = 0
Date  ADL

===============================================================

INVESTIGATIONS:

Date    Hb  TC  N L E MCV Plat  ESR RBS AC  PC  HbA1c Creat eGFR

  T3  T4  TSH TPOAb Iron  Ferritin  Transferrin TIBC

Date  Peripheral blood smear

Date  Specimen  Test

Date  BiliT SGOT  SGPT  GGT AP  Alb Glob  NH3 Uric Acid  Urine

Date  Lipid profile  CH/TG/HD/LD/VL Vitamin Na  K HCO3  Lactate  Ca  P Mg  PTH
    B12 D3

Date  PT / Prothrombin Time aPTT / Partial Thromboplastin Time  Thrombin Time    PSA CA-125
  Pt  Cont  INR Pt  Cont

Date  Trop-T  CPK LDH   Se Lactate  Se Pyruvate Se Homocysteine Se Ceruloplasmin  Plasma Procalcitonin

  CRP ANA-IF  ANA-Blot  RA Factor Anti-CCP  AChR Ab   Thyroid Ab (anti TPO)   Thyroglobulin Ab (TG Ab)

Date  Plasma Cortisol Post Synacthen  Prolactin PTH

Date  Test      Date  Test

Arterial Blood Gases
Date  pH  PCO2  pO2 sO2 Hb  Lactate Creat HCO3  Base Excess

CEREBRO-SPINAL FLUID
CSF WBC / cmm RBC Prot  Gluc  Pressure cm H2O
Date  TC  N%  L%  /cmm  mg/dl mg/dl Open  Close

Chest x-ray PA (   /  /   ):
Chest x-ray PA (   /  /   ):

X-ray    Spine (  /  /  ):
X-ray    Spine (  /  /  ):

ECG   (  /   /  ):           , QRS axis       deg.
ECG   (  /   /  ):           , QRS axis       deg.

24 hr ECG Holter   (  /   /  ):

ECHO (  /  / ):            , EF:    %     PASP:    mmHg.
ECHO (  /  / ):            , EF:    %     PASP:    mmHg.

TMT: (  /   /   ):
TMT: (  /   /   ):

Coronary Angiogram  (  /  /   ):

PFT / Pulmonary Function Test (  /  /  ):

NCS (  /   /   ):
NCS (  /   /   ):

EEG (  /   /   ):
EEG (  /   /   ):

CT Brain (  /   /   ):
CT Brain ( /   /  ):
CT Brain (  /   /   ):
CT Brain (  /   /   ):

MRI   Brain (  /   /  ):
MRI   Brain (  /   /  ):

MRI  Whole Spine (   /   /  ):
MRI  Whole Spine (  /   /  ):

4-vessel Duplex U/S scan (  /   /  ):
4-vessel Duplex U/S scan (  /   /  ):

U/S Abdomen & Pelvis (   /   /  ):
U/S Abdomen & Pelvis (   /   /  ):

U/S   -   Doppler   -  Limb (   /   /  ):

BMD Scan     (  /   /   ):  Total Body BMD T-score:

GI Endoscopy    (  /   /   ):

Audiogram  (  /   /   ):

Humphrey's Field Chart (   /  /  )

Sleep Study      (  /   /   ):
AHI:
RDI:
ODI:
Snore Index:        %

xx     (  /   /   ):

Biopsy (  /   /   ):

Psychological Assessment (  /  /  ):

Genomic Analysis  (  /  /  ):

REFERRALS:
 /  /  :

DISCUSSION:
 /  /  :

PLAN of Management:
 /  /  :

PATIENT EDUCATION & COUNSELING:
 /  /  :

CLINICAL COURSE & PROGRESS:

MEDICATIONS NOW:

MEDICATION  DOSE              TIME
            AM  Noon  PM  HS

PREVIOUS

Tab.                                INR
Date  Mon Tue Wed Thu Fri Sat   Sun
"""
        
        # Call GPT-4o-mini
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=6000
        )
        
        # Return plain text formatted report
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"❌ Diagnosis generation failed: {str(e)}")
        return None

def create_word_document(diagnosis_text, transcript=None):
    """
    Create a Word document from formatted diagnosis text
    
    Args:
        diagnosis_text (str): Pre-formatted medical report text
        transcript (str): Optional transcript text
        
    Returns:
        BytesIO: Word document as bytes
    """
    try:
        # Create document
        doc = Document()
        
        # Set narrow margins for medical report format
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add the pre-formatted diagnosis text directly
        # Use Courier New for monospaced format (tables align better)
        for line in diagnosis_text.split('\n'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        # Add transcript if provided
        if transcript:
            doc.add_page_break()
            doc.add_heading('CONSULTATION TRANSCRIPT', level=1)
            doc.add_paragraph(transcript)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
        
    except Exception as e:
        st.error(f"❌ Word document creation failed: {str(e)}")
        return None
